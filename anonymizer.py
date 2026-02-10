"""
Core DOCX Anonymization Module

This module handles the detection and anonymization of author-identifiable
information in DOCX files using a combination of:
- Named Entity Recognition (NER) for person names
- Keyword-based detection for affiliations, emails, etc.
- Positional analysis to identify author sections
- Hard constraint: Never modify content after References section
"""

from docx import Document
from docx.shared import RGBColor
from docx.oxml.xmlchemy import OxmlElement
import spacy
import re
from typing import List, Tuple, Set
import logging

logger = logging.getLogger(__name__)


class DocxAnonymizer:
    """
    Anonymizes author-identifiable information in DOCX files.
    
    Detection Strategy:
    1. Locate References/Bibliography section (hard boundary)
    2. Identify author section using keywords + position
    3. Use NER to detect person names
    4. Use patterns for emails, ORCIDs, affiliations
    5. Replace with placeholders while preserving formatting
    """
    
    # Keywords that indicate author-related sections (case-insensitive)
    AUTHOR_KEYWORDS = {
        'author', 'authors', 'affiliation', 'affiliations',
        'corresponding author', 'correspondence', 'email',
        'department', 'university', 'institution', 'institute'
    }
    
    # Section headers that mark the start of main content
    CONTENT_START_MARKERS = {
        'abstract', 'introduction', 'background', 
        'methods', 'methodology', 'materials and methods'
    }
    
    # Reference section markers (hard boundary - DO NOT cross)
    REFERENCE_MARKERS = {
        'references', 'bibliography', 'works cited',
        'literature cited', 'cited literature'
    }
    
    def __init__(self):
        """Initialize the anonymizer with NER model"""
        try:
            # Load English NER model
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded spaCy NER model successfully")
        except OSError:
            logger.error("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            raise
    
    def anonymize_document(self, input_path: str, output_path: str) -> bool:
        """
        Main entry point: anonymize a DOCX file.
        
        Args:
            input_path: Path to input DOCX file
            output_path: Path to save anonymized DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document(input_path)
            
            # Step 1: Find the References section boundary
            reference_start_idx = self._find_reference_section(doc)
            
            # Step 2: Find the author section boundary
            author_section_end = self._find_author_section_end(doc, reference_start_idx)
            
            # Step 3: Extract person names from author section using NER
            detected_names = self._extract_person_names(doc, author_section_end)
            
            # Step 4: Process paragraphs in author section only
            self._anonymize_author_section(doc, author_section_end, detected_names)
            
            # Save anonymized document
            doc.save(output_path)
            logger.info(f"Document anonymized successfully: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error anonymizing document: {str(e)}", exc_info=True)
            return False
    
    def _find_reference_section(self, doc: Document) -> int:
        """
        Find the start of the References/Bibliography section.
        
        This is a HARD BOUNDARY - we never modify content after this point.
        
        Args:
            doc: The document to search
            
        Returns:
            Index of first paragraph in References section, or len(paragraphs) if not found
        """
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            
            # Check if this is a reference section heading
            # Must be relatively short (likely a heading, not a sentence)
            if len(text) < 50 and any(marker in text for marker in self.REFERENCE_MARKERS):
                logger.info(f"Found References section at paragraph {idx}: '{para.text}'")
                return idx
        
        logger.warning("No References section found - will process entire document")
        return len(doc.paragraphs)
    
    def _find_author_section_end(self, doc: Document, reference_start: int) -> int:
        """
        Find where the author section ends (before main content).
        
        Author sections typically appear before Abstract or Introduction.
        
        Args:
            doc: The document to search
            reference_start: Index where References section begins
            
        Returns:
            Index of last paragraph in author section
        """
        # Search only up to References section
        for idx in range(min(reference_start, len(doc.paragraphs))):
            para = doc.paragraphs[idx]
            text = para.text.strip().lower()
            
            # Check for content start markers (Abstract, Introduction, etc.)
            if len(text) < 50 and any(marker in text for marker in self.CONTENT_START_MARKERS):
                logger.info(f"Found content start at paragraph {idx}: '{para.text}'")
                return idx
        
        # If no clear marker found, assume first 15 paragraphs as author section
        # (conservative estimate for most academic papers)
        default_end = min(15, reference_start)
        logger.info(f"No content marker found, using default author section end: {default_end}")
        return default_end
    
    def _extract_person_names(self, doc: Document, author_section_end: int) -> Set[str]:
        """
        Extract person names from author section using NER.
        
        Args:
            doc: The document to analyze
            author_section_end: Last paragraph index of author section
            
        Returns:
            Set of detected person names
        """
        names = set()
        
        # Combine text from author section
        author_text = ' '.join([
            para.text for para in doc.paragraphs[:author_section_end]
        ])
        
        # Run NER
        doc_nlp = self.nlp(author_text)
        
        # Extract PERSON entities
        for ent in doc_nlp.ents:
            if ent.label_ == 'PERSON':
                # Filter out single-letter names (likely initials in citations)
                if len(ent.text) > 2:
                    names.add(ent.text)
                    logger.debug(f"Detected person name: {ent.text}")
        
        logger.info(f"Extracted {len(names)} person names from author section")
        return names
    
    def _anonymize_author_section(self, doc: Document, author_section_end: int, 
                                   detected_names: Set[str]) -> None:
        """
        Anonymize author information in the author section.
        
        This is where the actual text replacement happens.
        
        Args:
            doc: The document to modify
            author_section_end: Last paragraph index of author section
            detected_names: Set of person names to anonymize
        """
        # Process only paragraphs in author section
        for idx in range(author_section_end):
            para = doc.paragraphs[idx]
            
            # Check each run in the paragraph (runs preserve formatting)
            for run in para.runs:
                original_text = run.text
                modified_text = original_text
                
                # 1. Anonymize person names
                for name in detected_names:
                    # Use word boundaries to avoid partial matches
                    pattern = r'\b' + re.escape(name) + r'\b'
                    modified_text = re.sub(
                        pattern, 
                        '[AUTHOR_NAME]', 
                        modified_text, 
                        flags=re.IGNORECASE
                    )
                
                # 2. Anonymize email addresses
                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                modified_text = re.sub(email_pattern, '[EMAIL]', modified_text)
                
                # 3. Anonymize ORCID IDs
                # Format: 0000-0002-1825-0097
                orcid_pattern = r'\b\d{4}-\d{4}-\d{4}-\d{3}[0-9X]\b'
                modified_text = re.sub(orcid_pattern, '[ORCID]', modified_text)
                
                # 4. Anonymize affiliation indicators
                # Look for patterns like "1University of X" or "Department of Y"
                if self._is_likely_affiliation(original_text):
                    # Check if it contains institutional keywords
                    if any(keyword in original_text.lower() 
                           for keyword in ['university', 'department', 'institute', 
                                         'college', 'school', 'center', 'centre']):
                        # Only replace if it's not just a mention in running text
                        # (i.e., appears to be an affiliation line)
                        if len(original_text.strip()) < 200 and not original_text.endswith('.'):
                            modified_text = '[AUTHOR_AFFILIATION]'
                
                # Update run text if modified
                if modified_text != original_text:
                    run.text = modified_text
                    logger.debug(f"Anonymized: '{original_text[:50]}...' -> '{modified_text[:50]}...'")
    
    def _is_likely_affiliation(self, text: str) -> bool:
        """
        Heuristic to determine if text is likely an affiliation line.
        
        Affiliations typically:
        - Start with a number or superscript marker
        - Are shorter than regular sentences
        - Contain institutional keywords
        
        Args:
            text: Text to analyze
            
        Returns:
            True if likely an affiliation
        """
        text_stripped = text.strip()
        
        # Empty or very short text
        if len(text_stripped) < 10:
            return False
        
        # Too long to be a typical affiliation line
        if len(text_stripped) > 300:
            return False
        
        # Starts with a digit (common affiliation marker)
        if text_stripped[0].isdigit():
            return True
        
        # Contains institutional keywords at the start
        first_words = ' '.join(text_stripped.split()[:3]).lower()
        if any(keyword in first_words 
               for keyword in ['department', 'university', 'institute', 'college']):
            return True
        
        return False


def test_anonymizer():
    """
    Test function to validate anonymization logic.
    """
    anonymizer = DocxAnonymizer()
    
    # Test NER
    test_text = "John Smith and Jane Doe from Harvard University."
    doc = anonymizer.nlp(test_text)
    print("\nNER Test:")
    for ent in doc.ents:
        print(f"  {ent.text} -> {ent.label_}")
    
    # Test email pattern
    test_emails = [
        "contact@university.edu",
        "john.smith@example.com",
        "invalid@email"
    ]
    print("\nEmail Pattern Test:")
    for email in test_emails:
        result = re.sub(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            '[EMAIL]',
            email
        )
        print(f"  {email} -> {result}")
    
    # Test ORCID pattern
    test_orcids = [
        "0000-0002-1825-0097",
        "0000-0001-5000-000X",
        "invalid-orcid"
    ]
    print("\nORCID Pattern Test:")
    for orcid in test_orcids:
        result = re.sub(
            r'\b\d{4}-\d{4}-\d{4}-\d{3}[0-9X]\b',
            '[ORCID]',
            orcid
        )
        print(f"  {orcid} -> {result}")


if __name__ == "__main__":
    test_anonymizer()
