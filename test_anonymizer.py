"""
Test script for DOCX Anonymization API

This script creates a sample document with author information
and tests the anonymization process.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from anonymizer import DocxAnonymizer
import os


def create_test_document(filename="test_paper.docx"):
    """
    Create a sample academic paper with author information.
    """
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Machine Learning Approaches to Climate Modeling")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    
    # Authors
    authors = doc.add_paragraph("John Smith¹, Jane Doe²*, Robert Johnson¹")
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.runs[0].font.size = Pt(12)
    
    # Affiliations
    affil1 = doc.add_paragraph("¹Department of Computer Science, Massachusetts Institute of Technology")
    affil1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil1.runs[0].font.size = Pt(10)
    
    affil2 = doc.add_paragraph("²Climate Research Institute, Harvard University")
    affil2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    affil2.runs[0].font.size = Pt(10)
    
    # Corresponding author
    corresp = doc.add_paragraph("*Corresponding author: jane.doe@harvard.edu")
    corresp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    corresp.runs[0].font.size = Pt(10)
    corresp.runs[0].italic = True
    
    # ORCID
    orcid = doc.add_paragraph("ORCID: 0000-0002-1825-0097")
    orcid.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    orcid.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Abstract
    abstract_heading = doc.add_paragraph("Abstract")
    abstract_heading.runs[0].bold = True
    abstract_heading.runs[0].font.size = Pt(14)
    
    abstract_text = doc.add_paragraph(
        "This paper presents a novel approach to climate modeling using machine learning. "
        "We demonstrate that our method achieves superior accuracy compared to traditional "
        "physical models while requiring significantly less computational resources. "
        "Our results show a 23% improvement in prediction accuracy for temperature forecasting."
    )
    abstract_text.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Introduction
    intro_heading = doc.add_paragraph("1. Introduction")
    intro_heading.runs[0].bold = True
    intro_heading.runs[0].font.size = Pt(14)
    
    intro_text = doc.add_paragraph(
        "Climate modeling has been a critical area of research for decades (Smith et al., 2020). "
        "Traditional approaches rely on complex physical simulations that require substantial "
        "computational power. Recent advances in machine learning offer promising alternatives "
        "that can complement or enhance traditional methods (Doe and Johnson, 2021)."
    )
    intro_text.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Methods
    methods_heading = doc.add_paragraph("2. Methods")
    methods_heading.runs[0].bold = True
    methods_heading.runs[0].font.size = Pt(14)
    
    methods_text = doc.add_paragraph(
        "We employed a deep neural network architecture with attention mechanisms. "
        "The model was trained on 50 years of historical climate data from the National "
        "Oceanic and Atmospheric Administration (NOAA). We used cross-validation to ensure "
        "robust performance across different geographical regions."
    )
    methods_text.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Results
    results_heading = doc.add_paragraph("3. Results")
    results_heading.runs[0].bold = True
    results_heading.runs[0].font.size = Pt(14)
    
    results_text = doc.add_paragraph(
        "Our model achieved an R² of 0.89 on the test set, significantly outperforming "
        "the baseline physical model (R² = 0.72). The improvement was particularly pronounced "
        "for short-term forecasts (1-7 days) where our model showed 31% better accuracy."
    )
    results_text.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # References
    ref_heading = doc.add_paragraph("References")
    ref_heading.runs[0].bold = True
    ref_heading.runs[0].font.size = Pt(14)
    
    # Important: These should NOT be anonymized
    ref1 = doc.add_paragraph(
        "Smith, J., Brown, A., & Wilson, C. (2020). Traditional climate modeling approaches. "
        "Journal of Climate Science, 45(3), 234-256."
    )
    ref1.runs[0].font.size = Pt(11)
    
    ref2 = doc.add_paragraph(
        "Doe, J., & Johnson, R. (2021). Machine learning in atmospheric science. "
        "Nature Climate Change, 11(2), 123-135."
    )
    ref2.runs[0].font.size = Pt(11)
    
    ref3 = doc.add_paragraph(
        "Thompson, M., Davis, K., & Martinez, L. (2019). Neural networks for weather prediction. "
        "AI in Earth Sciences, 8(4), 445-467."
    )
    ref3.runs[0].font.size = Pt(11)
    
    # Save document
    doc.save(filename)
    print(f"✓ Created test document: {filename}")
    return filename


def test_anonymization():
    """
    Test the anonymization process.
    """
    print("\n" + "="*60)
    print("DOCX ANONYMIZATION TEST")
    print("="*60 + "\n")
    
    # Create test document
    input_file = create_test_document()
    output_file = "test_paper_anonymized.docx"
    
    # Initialize anonymizer
    print("\n📝 Initializing anonymizer...")
    anonymizer = DocxAnonymizer()
    
    # Perform anonymization
    print(f"\n🔒 Anonymizing document: {input_file}")
    success = anonymizer.anonymize_document(input_file, output_file)
    
    if success:
        print(f"\n✅ SUCCESS! Anonymized document saved to: {output_file}")
        print("\n" + "="*60)
        print("VERIFICATION CHECKLIST")
        print("="*60)
        print("\nOpen both files and verify:")
        print("  ✓ Author names replaced with [AUTHOR_NAME]")
        print("  ✓ Affiliations replaced with [AUTHOR_AFFILIATION]")
        print("  ✓ Email replaced with [EMAIL]")
        print("  ✓ ORCID replaced with [ORCID]")
        print("  ✓ Abstract and Introduction are UNCHANGED")
        print("  ✓ References section is COMPLETELY UNCHANGED")
        print("  ✓ In-text citations (Smith et al., Doe and Johnson) are UNCHANGED")
        print("  ✓ Formatting (bold, italic, alignment) is preserved")
        print("\n" + "="*60 + "\n")
        
        # Display file sizes
        input_size = os.path.getsize(input_file)
        output_size = os.path.getsize(output_file)
        print(f"Input file size:  {input_size:,} bytes")
        print(f"Output file size: {output_size:,} bytes")
        print(f"Size difference:  {output_size - input_size:+,} bytes")
        
    else:
        print("\n❌ FAILED to anonymize document")
        return False
    
    return True


def test_edge_cases():
    """
    Test edge cases and boundary conditions.
    """
    print("\n" + "="*60)
    print("EDGE CASE TESTS")
    print("="*60 + "\n")
    
    # Test 1: Document without References section
    print("Test 1: Document without References section")
    doc = Document()
    doc.add_paragraph("Title: Test Paper")
    doc.add_paragraph("Author: Alice Johnson")
    doc.add_paragraph("Email: alice@university.edu")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This is the abstract text.")
    
    test1_input = "test_no_references.docx"
    test1_output = "test_no_references_anonymized.docx"
    doc.save(test1_input)
    
    anonymizer = DocxAnonymizer()
    result1 = anonymizer.anonymize_document(test1_input, test1_output)
    print(f"  Result: {'✓ PASS' if result1 else '✗ FAIL'}\n")
    
    # Test 2: Empty document
    print("Test 2: Empty document")
    doc2 = Document()
    doc2.add_paragraph("")
    
    test2_input = "test_empty.docx"
    test2_output = "test_empty_anonymized.docx"
    doc2.save(test2_input)
    
    result2 = anonymizer.anonymize_document(test2_input, test2_output)
    print(f"  Result: {'✓ PASS' if result2 else '✗ FAIL'}\n")
    
    # Test 3: Document with only References
    print("Test 3: Document with only References section")
    doc3 = Document()
    doc3.add_paragraph("References")
    doc3.add_paragraph("Smith, J. (2020). Paper title. Journal Name.")
    
    test3_input = "test_only_references.docx"
    test3_output = "test_only_references_anonymized.docx"
    doc3.save(test3_input)
    
    result3 = anonymizer.anonymize_document(test3_input, test3_output)
    print(f"  Result: {'✓ PASS' if result3 else '✗ FAIL'}\n")
    
    # Cleanup
    for f in [test1_input, test1_output, test2_input, test2_output, 
              test3_input, test3_output]:
        if os.path.exists(f):
            os.remove(f)
    
    print("="*60 + "\n")


if __name__ == "__main__":
    # Run main test
    test_anonymization()
    
    # Run edge case tests
    test_edge_cases()
    
    print("\n✨ All tests completed!\n")
